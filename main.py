from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response, HTMLResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel
from typing import Optional
from io import BytesIO
from pathlib import Path
import uuid
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from scenarios.receipt_simple import ReceiptSimpleScenario
from scenarios.receipt_advanced import ReceiptAdvancedScenario
from scenarios.loan import LoanScenario
from scenarios.claim_marketplace_buyer import ClaimMarketplaceBuyerScenario

app = FastAPI()

BASE_DIR = Path(__file__).resolve().parent

# Подключаем static файлы если папка существует
static_dir = BASE_DIR / "static"
if static_dir.exists():
    app.mount("/static", StaticFiles(directory=str(static_dir)), name="static")

# Маршрут для главной страницы
@app.get("/", response_class=HTMLResponse)
async def root():
    index_path = BASE_DIR / "index.html"
    if index_path.exists():
        return HTMLResponse(index_path.read_text(encoding="utf-8"))
    return HTMLResponse("<html><body><h1>ProjectText</h1><p>index.html not found</p></body></html>")

# Настройка CORS - разрешаем все источники
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Хранилище сессий: ключ = f"{session_id}:{scenario_type}"
sessions = {}

# Словарь шаблонов для каждого типа сценария
template_map = {
    "receipt_simple": "templates/receipt_simple.txt",
    "receipt_advanced": "templates/receipt_advanced.txt",
    "loan": "templates/loan.txt",
    "claim_marketplace_buyer": "templates/claim_marketplace_buyer.txt"
}

class AgentResponse(BaseModel):
    is_complete: bool = False
    question: Optional[str] = None
    document: Optional[str] = None
    error: Optional[str] = None
    session_id: str
    current_step: Optional[str] = None
    optional: bool = False
    field_type: Optional[str] = None
    current_value: Optional[str] = None
    collected_data: dict = {}

class ScenarioRequest(BaseModel):
    session_id: str
    answer: str = ""

class DocxRequest(BaseModel):
    text: str
    scenario_type: Optional[str] = None
    collected_data: Optional[dict] = None

def generate_docx_filename(scenario_type: str = None, collected_data: dict = None) -> str:
    """Генерирует осмысленное имя файла на основе типа документа и даты."""
    from datetime import datetime
    
    date_str = None
    if collected_data:
        date_value = collected_data.get('date')
        if date_value:
            try:
                if isinstance(date_value, str):
                    if '.' in date_value:
                        parts = date_value.split('.')
                        if len(parts) == 3:
                            date_str = f"{parts[0]}.{parts[1]}.{parts[2]}"
                        else:
                            date_obj = datetime.strptime(date_value, '%Y-%m-%d')
                            date_str = date_obj.strftime('%d.%m.%Y')
                    else:
                        date_obj = datetime.strptime(str(date_value)[:10], '%Y-%m-%d')
                        date_str = date_obj.strftime('%d.%m.%Y')
                else:
                    date_obj = datetime.strptime(str(date_value)[:10], '%Y-%m-%d')
                    date_str = date_obj.strftime('%d.%m.%Y')
            except:
                date_str = datetime.now().strftime('%d.%m.%Y')
        else:
            date_str = datetime.now().strftime('%d.%m.%Y')
    else:
        date_str = datetime.now().strftime('%d.%m.%Y')
    
    name_map = {
        'loan': 'Договор_займа',
        'receipt_simple': 'Расписка_простая',
        'receipt_advanced': 'Расписка_расширенная',
        'claim_marketplace_buyer': 'Претензия_маркетплейс'
    }
    prefix = name_map.get(scenario_type, 'Документ') if scenario_type else 'Документ'
    
    return f"{prefix}_{date_str}.docx"

def format_signatures(document_text: str) -> str:
    """Приводит подписи к единому читаемому формату."""
    import re
    
    doc_text = document_text
    
    doc_text = re.sub(r'_{3,}\s*/\s*([А-ЯЁа-яёё\s\-\.]+)', r'______________ / \1', doc_text)
    
    has_lender = bool(re.search(r'Подпись.*займодавца', doc_text, re.IGNORECASE))
    has_borrower = bool(re.search(r'Подпись.*заемщика', doc_text, re.IGNORECASE))
    has_single_sig = bool(re.search(r'Подпись.*:?[\s_]*$', doc_text, re.IGNORECASE | re.MULTILINE))
    
    if has_lender and has_borrower:
        if '____________ / {lender}' in doc_text or '_____________ / {lender}' in doc_text:
            doc_text = re.sub(r'____________+ / \{lender\}', '______________ / ФИО займодавца', doc_text)
        if '____________ / {borrower}' in doc_text or '_____________ / {borrower}' in doc_text:
            doc_text = re.sub(r'____________+ / \{borrower\}', '______________ / ФИО заемщика', doc_text)
    elif has_single_sig:
        if '____________ / {fio_receiver}' in doc_text or '_____________ / {fio_receiver}' in doc_text:
            doc_text = re.sub(r'____________+ / \{fio_receiver\}', '______________ / ФИО', doc_text)
    
    if '______________ / ФИО' in doc_text or '____________ / ФИО' in doc_text:
        pass
    
    return doc_text

def generate_docx(document_text: str) -> bytes:
    """Генерирует красивый DOCX документ из текста."""
    import re
    document_text = format_signatures(document_text)
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    for paragraph in doc.paragraphs:
        paragraph.paragraph_format.line_spacing = 1.5
    
    doc.add_paragraph()
    
    table_pattern = re.compile(r'<div class="schedule-table"><table>.*?</table></div>', re.DOTALL)
    tables_html = {}
    for i, match in enumerate(table_pattern.findall(document_text)):
        tables_html[f'__TABLE_{i}__'] = match
        document_text = document_text.replace(match, f'__TABLE_{i}__')
    
    lines = document_text.split('\n')
    prev_was_empty = True
    
    for line in lines:
        line = line.strip()
        if not line:
            if not prev_was_empty:
                prev_was_empty = True
            continue
        
        prev_was_empty = False
        
        if line.startswith('__TABLE_') and line.endswith('__'):
            table_html = tables_html.get(line)
            if table_html:
                tr_pattern = re.compile(r'<tr>.*?</tr>', re.DOTALL)
                th_pattern = re.compile(r'<th>(.*?)</th>', re.DOTALL)
                td_pattern = re.compile(r'<td[^>]*>(.*?)</td>', re.DOTALL)
                
                rows = tr_pattern.findall(table_html)
                if rows:
                    table = doc.add_table(rows=len(rows), cols=6)
                    table.style = 'Table Grid'
                    
                    for row_idx, row_html in enumerate(rows):
                        cells = td_pattern.findall(row_html)
                        if row_idx == 0 and th_pattern.findall(row_html):
                            headers = th_pattern.findall(row_html)
                            for col_idx, header_text in enumerate(headers):
                                if col_idx < 6:
                                    cell = table.rows[row_idx].cells[col_idx]
                                    cell.text = header_text
                                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for run in cell.paragraphs[0].runs:
                                        run.font.bold = True
                                        run.font.name = 'Times New Roman'
                        else:
                            cell_texts = td_pattern.findall(row_html)
                            for col_idx, cell_text in enumerate(cell_texts):
                                if col_idx < 6:
                                    cell = table.rows[row_idx].cells[col_idx]
                                    cell.text = cell_text
                                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                                    for run in cell.paragraphs[0].runs:
                                        run.font.name = 'Times New Roman'
                continue
        
        is_header = (
            line.isupper() or
            'РАСПИСКА' in line or
            'ДОГОВОР' in line or
            'ИСКОВОЕ ЗАЯВЛЕНИЕ' in line or
            'ГРАФИК ПЛАТЕЖЕЙ' in line or
            'ПРАВОВАЯ ИНФОРМАЦИЯ' in line
        )
        
        if is_header:
            p = doc.add_paragraph(line)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(6)
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(14)
                run.font.name = 'Times New Roman'
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Cm(1.25)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.5
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    
    doc.add_paragraph()
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

@app.post("/download-docx")
def download_docx(request: DocxRequest):
    """Скачивание документа в формате DOCX."""
    docx_bytes = generate_docx(request.text)
    filename = generate_docx_filename(request.scenario_type, request.collected_data)
    return Response(
        content=docx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )

def _make_key(session_id: str, scenario_type: str) -> str:
    """Создаёт составной ключ для хранения сценария."""
    return f"{session_id}:{scenario_type}"

def get_or_create_scenario(session_id: str = None, scenario_type: str = "receipt_simple"):
    """Получает или создаёт сценарий по типу."""
    # Если session_id не передан, создаём новый
    if session_id is None or session_id == "":
        session_id = str(uuid.uuid4())
    
    key = _make_key(session_id, scenario_type)
    
    if key not in sessions:
        if scenario_type == "receipt_advanced":
            sessions[key] = ReceiptAdvancedScenario()
        elif scenario_type == "loan":
            sessions[key] = LoanScenario()
        elif scenario_type == "claim_marketplace_buyer":
            sessions[key] = ClaimMarketplaceBuyerScenario()
        elif scenario_type == "claim_simple":
            # Заглушка для claim_simple - возвращаем None
            sessions[key] = None
        else:
            sessions[key] = ReceiptSimpleScenario()
    
    return sessions[key], session_id

def is_error_response(text: str) -> bool:
    """Проверяет, является ли текст сообщением об ошибке валидации."""
    if not text:
        return False
    error_prefixes = ("не может быть", "Пожалуйста, введите", "Ошибка")
    return text.startswith(error_prefixes)

@app.post("/api/scenario/{scenario_type}", response_model=AgentResponse)
def handle_scenario(request: ScenarioRequest, scenario_type: str):
    """Универсальный эндпоинт для всех типов сценариев."""
    # Проверяем, поддерживается ли тип
    if scenario_type == "claim_simple":
        return AgentResponse(
            session_id=request.session_id or str(uuid.uuid4()),
            error="Сценарий в разработке",
            current_step="unavailable",
            collected_data={}
        )
    
    # Получаем или создаём сценарий
    scenario, session_id = get_or_create_scenario(request.session_id, scenario_type)
    
    # Проверяем, реализован ли сценарий
    if scenario is None:
        return AgentResponse(
            session_id=session_id,
            error="Сценарий в разработке",
            current_step="unavailable",
            collected_data={}
        )
    
    template_path = template_map.get(scenario_type, "templates/receipt_simple.txt")
    
    # Если есть ответ - обрабатываем его (включая пустую строку для пропуска optional полей)
    if request.answer is not None:
        scenario.process_answer(request.answer)
        
        # Сразу проверяем завершение
        if scenario.is_complete():
            document = scenario.generate_document(template_path)
            return AgentResponse(
                is_complete=True,
                document=document,
                session_id=session_id,
                current_step="done",
                collected_data=scenario.data
            )
        
        # Если не завершён, получаем следующий вопрос
        next_question = scenario.get_next_question()
        cv = None
        if getattr(scenario, '_return_to_preview', False):
            idx = scenario._current_index
            if 0 <= idx < len(scenario._steps):
                step = scenario._steps[idx]
                if step.data_key and step.data_key in scenario.data:
                    cv = str(scenario.data[step.data_key])
        return AgentResponse(
            question=next_question,
            session_id=session_id,
            current_step=scenario.get_current_step(),
            optional=getattr(scenario, 'is_current_optional', lambda: False)(),
            field_type=getattr(scenario, 'get_current_field_type', lambda: None)(),
            current_value=cv,
            collected_data=scenario.data
        )
    
    # Первый вызов без ответа - инициализируем сценарий
    if scenario.get_current_step() == "start":
        scenario.process_answer("")
    
    # Получаем первый вопрос
    question = scenario.get_next_question()
    
    return AgentResponse(
        question=question,
        session_id=session_id,
        current_step=scenario.get_current_step(),
        optional=getattr(scenario, 'is_current_optional', lambda: False)(),
        field_type=getattr(scenario, 'get_current_field_type', lambda: None)(),
        collected_data=scenario.data
    )

@app.post("/api/session/{session_id}/reset")
def reset_session(session_id: str):
    """Сбрасывает все сценарии для данного session_id."""
    # Удаляем все сессии для данного session_id (все типы)
    keys_to_delete = [key for key in sessions if key.startswith(f"{session_id}:")]
    for key in keys_to_delete:
        del sessions[key]
    
    return {"status": "ok", "session_id": session_id, "deleted": len(keys_to_delete)}

@app.get("/api/session/{session_id}/status")
def session_status(session_id: str):
    """Возвращает статус всех сценариев для данного session_id."""
    result = {"session_id": session_id, "scenarios": {}}
    
    for key, scenario in sessions.items():
        if key.startswith(f"{session_id}:"):
            scenario_type = key.split(":")[1]
            if scenario is not None:
                result["scenarios"][scenario_type] = {
                    "current_step": scenario.get_current_step(),
                    "is_complete": scenario.is_complete(),
                    "data": scenario.data
                }
            else:
                result["scenarios"][scenario_type] = {"error": "not_implemented"}
    
    if not result["scenarios"]:
        return {"status": "not_found", "session_id": session_id}
    
    return result

# Обратная совместимость - старые эндпоинты
@app.post("/api/scenario/receipt_simple", response_model=AgentResponse)
def receipt_simple(request: ScenarioRequest):
    return handle_scenario(request, "receipt_simple")

@app.post("/api/scenario/receipt_advanced", response_model=AgentResponse)
def receipt_advanced(request: ScenarioRequest):
    return handle_scenario(request, "receipt_advanced")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)